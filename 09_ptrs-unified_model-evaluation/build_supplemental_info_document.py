"""Build Document S1 — the supplemental-information compendium.

Contents:
  - Cover page
  - "Key" / table of contents listing Supplementary Note + Figures S1–S6 + Tables S1–S7
  - Supplementary Note (full text preserved from Supplementary_Note_AUC_ceiling.docx)
  - Supplemental references section (placeholder — copy from manuscript when ready)

Outputs (written to both):
  - <repo>/09_ptrs-unified_model-evaluation/supplement/Document_S1_Supplemental_Information.docx
  - /Users/nancyh/Desktop/Document_S1_Supplemental_Information.docx
"""
import warnings; warnings.filterwarnings('ignore')
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

DESKTOP = Path('/Users/nancyh/Desktop')
REPO    = Path('/Users/nancyh/Desktop/asthma-prs-study-fresh/09_ptrs-unified_model-evaluation')
SUPP    = REPO / 'supplement'; SUPP.mkdir(exist_ok=True)
OUT_TARGETS = [SUPP / 'Document_S1_Supplemental_Information.docx',
               DESKTOP / 'Document_S1_Supplemental_Information.docx']

# =====================================================================
# Key entries (paper's structure)
# =====================================================================
KEY_ITEMS = [
    ('Supplementary Note', [
        ('Theoretical common-variant AUC ceiling',
         'Liability-threshold–framework derivation of the maximum AUC attainable by a predictor '
         'that perfectly captures the common-variant genetic liability to asthma. '
         'Reports point estimate and a 95% confidence interval propagated from LDSC heritability.'),
    ]),
    ('Figures', [
        ('Figure S1', 'FUSION TWAS Miami plots of gene–tissue associations across 39 bulk GTEx tissues.'),
        ('Figure S2', 'FUSION TWAS Miami plots of gene–cell-type associations across 17 pseudobulked OneK1K PBMC cell types.'),
        ('Figure S3', 'FUSION TWAS associations tested versus passing nominal filters across GTEx tissues and OneK1K cell types.'),
        ('Figure S4', 'GO term enrichment of FDR-significant genes from FUSION TWAS gene–tissue and gene–cell-type associations.'),
        ('Figure S5', 'Number of genes in the FOCUS 90% credible set per GTEx tissue and OneK1K cell type.'),
        ('Figure S6', 'FOCUS posterior support is more concentrated in OneK1K cell types than in GTEx tissues.'),
    ]),
    ('Tables', [
        ('Table S1', 'European-ancestry meta-analysis index variants — 275 independent GWS loci (P ≤ 5 × 10⁻⁸) from the TAGC + GBMI meta-analysis. Unchanged content.'),
        ('Table S2', 'Gene–context pairs retained at each stage of TWAS and FOCUS fine-mapping — 3-row funnel across GTEx bulk tissue and OneK1K cell types. Unchanged content.'),
        ('Table S3', 'Per-feature PTRS models retained by the cross-cohort consistency filter — TWAS P+T shortlist across all 4 candidate p-value thresholds × 2 modalities.'),
        ('Table S4', 'Performance of tissue-level PTRS from the TWAS P+T pipeline (all p-value thresholds) and the MA-FOCUS pipeline — all classifiers. Two-block table distinguished by a Pipeline column.'),
        ('Table S5', 'Performance of cell-type-level PTRS from the TWAS P+T pipeline (all p-value thresholds) and the MA-FOCUS pipeline — all classifiers. Same two-block layout as Table S4.'),
        ('Table S6', 'Feature-count ablation for the cross-modal PTRS + PRS integration — 2 anchor baselines + 18 three-feature ablations × 2 PRS variants. All 18/18 negative, mean ΔAUC = −0.031.'),
        ('Table S7', 'Best-performing models across scoring categories — one row per (PRS baseline, TWAS P+T single-feature, MA-FOCUS single-feature, unified PTRS per modality, cross-modal integrated per PRS). Ranked by CAMP-Balanced AUC.'),
    ]),
    ('Supplemental references', [
        ('References list',
         'Bibliographic references cited in the supplementary note and captions. '
         'To be finalized alongside the main manuscript reference list.'),
    ]),
]

# =====================================================================
# Build the docx
# =====================================================================
doc = Document()

# --- Page style: single column, 11 pt body ---
styles = doc.styles
n_style = styles['Normal']
n_style.font.name = 'Calibri'
n_style.font.size = Pt(11)

# --- Title page ---
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Supplemental Information')
run.bold = True
run.font.size = Pt(20)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Document S1. Supplementary Note, Figures S1–S6, Tables S1–S7, and supplemental references.')
run.italic = True
run.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Multimodal causal gene prioritization and nonlinear tissue and cell-type '
                'interaction risk modeling reveal the regulatory architecture of childhood-onset asthma.')
run.font.size = Pt(11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Nan Huang*, Michelle F. Ragsac*, Xiaoyu Gui, Kelan G. Tantisira, Tiffany Amariuta⛒')
run.font.size = Pt(10)

doc.add_paragraph()  # spacer

# --- KEY / Table of contents ---
h = doc.add_paragraph()
run = h.add_run('Key (contents of this supplement)')
run.bold = True
run.font.size = Pt(15)

for section_title, items in KEY_ITEMS:
    ph = doc.add_paragraph()
    r = ph.add_run(section_title)
    r.bold = True
    r.font.size = Pt(13)
    for label, description in items:
        p = doc.add_paragraph(style='List Bullet')
        rlabel = p.add_run(f'{label}. ')
        rlabel.bold = True
        rdesc = p.add_run(description)

doc.add_page_break()

# --- Supplementary Note (verbatim from AUC-ceiling doc) ---
ph = doc.add_paragraph()
r = ph.add_run('Supplementary Note — Theoretical common-variant AUC ceiling')
r.bold = True
r.font.size = Pt(15)

note_src = DESKTOP / 'Supplementary_Note_AUC_ceiling.docx'
if note_src.exists():
    src = Document(note_src)
    for para in src.paragraphs:
        text = para.text.strip()
        if not text:
            doc.add_paragraph()
            continue
        # Skip the title (already added)
        if text.startswith('Supplementary Note') and 'AUC ceiling' in text:
            continue
        p = doc.add_paragraph()
        p.add_run(text)
else:
    doc.add_paragraph('[Supplementary_Note_AUC_ceiling.docx not found on Desktop — copy the note here.]')

doc.add_page_break()

# --- Figures pointer ---
ph = doc.add_paragraph()
r = ph.add_run('Figures S1–S6')
r.bold = True
r.font.size = Pt(15)
doc.add_paragraph(
    'The figure PDFs / PNGs for Figures S1–S6 are provided as separate image files '
    'alongside this document. See the Key above for the one-line description of each figure. '
    'Full legends appear on the corresponding figure pages.')

doc.add_page_break()

# --- Tables pointer ---
ph = doc.add_paragraph()
r = ph.add_run('Tables S1–S7')
r.bold = True
r.font.size = Pt(15)
doc.add_paragraph(
    'Supplementary Tables S1–S7 are provided as an Excel workbook '
    '(Supplementary_Tables_PTRS_fresh.xlsx) plus one standalone workbook for the '
    'cross-modal feature-count ablation (Supplementary_Table_Feature_Ablation_fresh.xlsx). '
    'Each sheet is prefixed with a bold title row and a wrapped caption row above the data '
    'header. See the Key above for one-line summaries.')

doc.add_page_break()

# --- Supplemental references ---
ph = doc.add_paragraph()
r = ph.add_run('Supplemental references')
r.bold = True
r.font.size = Pt(15)
doc.add_paragraph(
    '[To be finalized. Copy the supplemental references block from the main manuscript here.]')

for target in OUT_TARGETS:
    doc.save(target)
    print(f'wrote {target}')
